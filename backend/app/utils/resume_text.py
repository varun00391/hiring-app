"""Extract plain text from PDF and DOCX for resume parsing.

PDF:
- Primary: PyMuPDF — page text + table detection (vector tables).
- Fallback: pypdf text extraction when PyMuPDF yields nothing.
- Scanned / image-only PDFs: render pages and OCR with Tesseract when text is too thin.

DOCX:
- Walks document body in order: paragraphs and tables (including nested tables).
- Includes header/footer paragraph text.

Requires system `tesseract-ocr` for OCR (see Dockerfile)."""

from __future__ import annotations

import asyncio
import io
import logging

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

# Optional: PyMuPDF (PDF text + tables + render for OCR)
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None  # type: ignore[misc, assignment]

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None  # type: ignore[misc, assignment]
    Image = None  # type: ignore[misc, assignment]

logger = logging.getLogger(__name__)


class ResumeTextError(Exception):
    pass


# If extracted text is shorter than this (after strip), treat as likely scanned and run OCR.
_MIN_CHARS_BEFORE_OCR = 50
# Cap OCR work for large resume PDFs
_MAX_OCR_PAGES = 30
_OCR_ZOOM = 2.0


def _pdf_text_pypdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _fitz_page_table_lines(page) -> list[str]:
    """Extract grid rows from vector tables on a PDF page."""
    out: list[str] = []
    if fitz is None:
        return out
    try:
        finder = page.find_tables()
    except Exception as exc:
        logger.debug("PDF page find_tables: %s", exc)
        return out
    tables = getattr(finder, "tables", None) or []
    for tab in tables:
        try:
            rows = tab.extract()
        except Exception as exc:
            logger.debug("table.extract: %s", exc)
            continue
        if not rows:
            continue
        for row in rows:
            if not row:
                continue
            cells = [str(c).strip() if c is not None else "" for c in row]
            line = " | ".join(cells)
            if line.strip():
                out.append(line)
    return out


def _pdf_fitz_text_and_tables(data: bytes) -> str:
    if fitz is None:
        return ""
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        parts: list[str] = []
        for page in doc:
            txt = (page.get_text("text") or "").strip()
            if txt:
                parts.append(txt)
            parts.extend(_fitz_page_table_lines(page))
        return "\n\n".join(parts)
    finally:
        doc.close()


def _normalized_tokens(s: str) -> str:
    return " ".join(s.lower().split())


def _ocr_likely_redundant(native: str, ocr: str) -> bool:
    """If OCR repeats what we already got from text layers, skip merging OCR output."""
    if not native.strip() or not ocr.strip():
        return False
    n, o = _normalized_tokens(native), _normalized_tokens(ocr)
    if n == o:
        return True
    if len(n) >= 12 and (n in o or o in n):
        return True
    return False


def _pdf_ocr(data: bytes) -> str:
    if fitz is None:
        raise ResumeTextError("PyMuPDF is not installed (required for scanned PDF support)")
    if pytesseract is None or Image is None:
        raise ResumeTextError("pytesseract/Pillow not installed")

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        n = min(len(doc), _MAX_OCR_PAGES)
        chunks: list[str] = []
        mat = fitz.Matrix(_OCR_ZOOM, _OCR_ZOOM)
        for i in range(n):
            page = doc[i]
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            try:
                chunk = pytesseract.image_to_string(img, lang="eng")
            except pytesseract.TesseractNotFoundError as exc:
                raise ResumeTextError(
                    "Tesseract OCR is not installed on the server. "
                    "Install the `tesseract-ocr` package (e.g. apt install tesseract-ocr)."
                ) from exc
            except Exception as exc:
                raise ResumeTextError(f"OCR failed: {exc}") from exc
            if chunk and chunk.strip():
                chunks.append(chunk.strip())
        return "\n\n".join(chunks)
    finally:
        doc.close()


def _extract_pdf_sync(data: bytes) -> str:
    """Full PDF pipeline: native text + tables; OCR when native text is insufficient."""
    if fitz is not None:
        merged = _pdf_fitz_text_and_tables(data).strip()
    else:
        merged = ""

    if not merged:
        merged = _pdf_text_pypdf(data).strip()

    if len(merged) >= _MIN_CHARS_BEFORE_OCR:
        return merged

    # Short or empty — try OCR (scanned PDFs)
    try:
        ocr_text = _pdf_ocr(data).strip()
    except ResumeTextError:
        if merged:
            logger.warning("OCR unavailable or failed; using partial native PDF text only")
            return merged
        raise

    if ocr_text:
        if merged and _ocr_likely_redundant(merged, ocr_text):
            logger.debug("PDF: OCR output redundant with native text; keeping native only")
            return merged
        if merged:
            logger.info(
                "PDF: combined native short text with OCR (%d chars native, %d OCR)",
                len(merged),
                len(ocr_text),
            )
            return f"{merged}\n\n--- OCR ---\n\n{ocr_text}"
        logger.info("PDF: OCR-only extraction (%d chars)", len(ocr_text))
        return ocr_text

    if merged:
        return merged

    raise ResumeTextError(
        "No text could be extracted from PDF. If this is a scanned file, ensure Tesseract OCR is installed."
    )


async def pdf_to_text(data: bytes) -> str:
    try:
        return await asyncio.to_thread(_extract_pdf_sync, data)
    except ResumeTextError:
        raise
    except Exception as exc:
        raise ResumeTextError(str(exc)) from exc


def _extract_table_rows(table: Table) -> list[str]:
    """Flatten a table to text lines; handles nested tables inside cells."""
    lines: list[str] = []
    for row in table.rows:
        cells_text: list[str] = []
        for cell in row.cells:
            bits: list[str] = []
            for p in cell.paragraphs:
                t = p.text.strip()
                if t:
                    bits.append(t)
            for inner in cell.tables:
                inner_rows = _extract_table_rows(inner)
                if inner_rows:
                    bits.append(" / ".join(inner_rows))
            cells_text.append(" ".join(bits).strip())
        line = " | ".join(cells_text)
        if line.strip():
            lines.append(line)
    return lines


def _docx_headers_footers(document: Document) -> list[str]:
    out: list[str] = []
    for section in document.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                t = p.text.strip()
                if t:
                    out.append(t)
    return out


def _docx_body_in_order(document: Document) -> list[str]:
    lines: list[str] = []
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            para = Paragraph(child, document)
            t = para.text.strip()
            if t:
                lines.append(t)
        elif tag.endswith("}tbl"):
            tbl = Table(child, document)
            lines.extend(_extract_table_rows(tbl))
    return lines


def _extract_docx_sync(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    header_footer = _docx_headers_footers(document)
    body = _docx_body_in_order(document)
    chunks = []
    if header_footer:
        chunks.append("\n".join(header_footer))
    if body:
        chunks.append("\n".join(body))
    merged = "\n\n".join(chunks).strip()
    if not merged:
        raise ResumeTextError("No text extracted from DOCX")
    return merged


async def docx_to_text(data: bytes) -> str:
    try:
        return await asyncio.to_thread(_extract_docx_sync, data)
    except ResumeTextError:
        raise
    except Exception as exc:
        raise ResumeTextError(str(exc)) from exc
