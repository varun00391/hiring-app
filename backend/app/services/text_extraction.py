"""Text extraction from PDF, DOC, and DOCX files."""

import asyncio
import io
import logging
import shutil
import subprocess
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from app.exceptions import CorruptedFileError, ExtractionError
from app.models.schemas import FileType

logger = logging.getLogger(__name__)


class TextExtractor:
    """Async wrapper around synchronous text extraction libraries."""

    async def extract(self, file_bytes: bytes, file_type: FileType, filename: str) -> tuple[str, str]:
        """Extract text and return (text, extraction_method)."""
        try:
            if file_type == FileType.PDF:
                return await asyncio.to_thread(self._extract_pdf, file_bytes)
            if file_type == FileType.DOCX:
                return await asyncio.to_thread(self._extract_docx, file_bytes)
            if file_type == FileType.DOC:
                return await asyncio.to_thread(self._extract_doc, file_bytes, filename)
        except (CorruptedFileError, ExtractionError):
            raise
        except Exception as exc:
            logger.exception("Unexpected extraction failure")
            raise ExtractionError(
                "Failed to extract text from the uploaded file.",
                details={"file_type": file_type.value, "reason": str(exc)},
            ) from exc

        raise ExtractionError("Unsupported file type for extraction.")

    def _extract_pdf(self, file_bytes: bytes) -> tuple[str, str]:
        text_parts: list[str] = []
        method = "pdfplumber"

        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
        except Exception as exc:
            logger.warning("pdfplumber extraction failed, falling back to PyMuPDF", exc_info=exc)
            text_parts = []
            method = "pymupdf"

        if not text_parts or sum(len(t) for t in text_parts) < 50:
            try:
                text_parts = []
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    for page in doc:
                        page_text = page.get_text("text") or ""
                        if page_text.strip():
                            text_parts.append(page_text)
                method = "pymupdf" if method == "pymupdf" else "pdfplumber+pymupdf"
            except Exception as exc:
                raise CorruptedFileError(
                    "The PDF file appears to be corrupted or unreadable.",
                    details={"reason": str(exc)},
                ) from exc

        text = "\n\n".join(text_parts).strip()
        if not text:
            raise ExtractionError(
                "No text could be extracted from the PDF. The file may be image-only or scanned.",
                details={"hint": "Consider OCR preprocessing for scanned resumes."},
            )
        return text, method

    def _extract_docx(self, file_bytes: bytes) -> tuple[str, str]:
        try:
            document = Document(io.BytesIO(file_bytes))
        except Exception as exc:
            raise CorruptedFileError(
                "The DOCX file appears to be corrupted or unreadable.",
                details={"reason": str(exc)},
            ) from exc

        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))

        text = "\n".join(paragraphs + table_text).strip()
        if not text:
            raise ExtractionError("No text could be extracted from the DOCX file.")
        return text, "python-docx"

    def _extract_doc(self, file_bytes: bytes, filename: str) -> tuple[str, str]:
        """Extract legacy .doc files using LibreOffice conversion when available."""
        libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not libreoffice:
            raise ExtractionError(
                "Legacy .doc extraction requires LibreOffice (soffice) in the runtime environment.",
                details={"filename": filename},
            )

        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / Path(filename).name
            input_path.write_bytes(file_bytes)

            try:
                subprocess.run(
                    [
                        libreoffice,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        tmpdir,
                        str(input_path),
                    ],
                    check=True,
                    capture_output=True,
                    timeout=60,
                )
            except subprocess.CalledProcessError as exc:
                raise CorruptedFileError(
                    "Failed to convert .doc file. The file may be corrupted.",
                    details={"stderr": exc.stderr.decode(errors="ignore") if exc.stderr else ""},
                ) from exc
            except subprocess.TimeoutExpired as exc:
                raise ExtractionError("Timed out while converting .doc file.") from exc

            output_path = input_path.with_suffix(".txt")
            if not output_path.exists():
                raise ExtractionError("LibreOffice did not produce a text output for the .doc file.")

            text = output_path.read_text(encoding="utf-8", errors="ignore").strip()
            if not text:
                raise ExtractionError("No text could be extracted from the .doc file.")
            return text, "libreoffice-doc-to-txt"
